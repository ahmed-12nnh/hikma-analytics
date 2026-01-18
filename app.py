import streamlit as st
import google.generativeai as genai
import PyPDF2
import pandas as pd
from io import StringIO, BytesIO
import base64
import re
from collections import Counter

# ---------------------------------------------------------
# 📦 استيراد المكتبات الإضافية
# ---------------------------------------------------------
try:
    from weasyprint import HTML as WeasyHTML
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ---------------------------------------------------------
# 🔑 إعدادات المفتاح
# ---------------------------------------------------------
try:
    API_KEY = st.secrets["GOOGLE_API_KEY"]
except:
    API_KEY = None

# ---------------------------------------------------------
# 🎨 إعدادات الصفحة
# ---------------------------------------------------------
st.set_page_config(
    page_title="منصة التحليل الاستراتيجي",
    page_icon="🦅",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ---------------------------------------------------------
# 🎨 CSS للواجهة الرئيسية
# ---------------------------------------------------------
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Tajawal:wght@300;400;500;700;800;900&display=swap');
    
    * { box-sizing: border-box; }
    
    .stApp {
        background: linear-gradient(135deg, #0a1628 0%, #1a365d 50%, #0a1628 100%);
        font-family: 'Tajawal', sans-serif;
        direction: rtl;
    }

    [data-testid="stSidebar"] { display: none; }
    header { visibility: hidden; }
    #MainMenu { visibility: hidden; }
    footer { visibility: hidden; }
    [data-testid="stToolbar"] { display: none; }

    .hero-section {
        background: linear-gradient(135deg, rgba(26, 54, 93, 0.95), rgba(45, 74, 113, 0.9));
        border-radius: 20px;
        padding: 50px 30px;
        text-align: center;
        margin: 20px;
        border: 2px solid rgba(212, 175, 55, 0.4);
        box-shadow: 0 20px 60px rgba(0, 0, 0, 0.4);
        position: relative;
    }

    .hero-section::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        height: 4px;
        background: linear-gradient(90deg, transparent, #d4af37, transparent);
    }

    .main-title {
        font-size: 48px;
        font-weight: 900;
        background: linear-gradient(180deg, #d4af37 0%, #f4d03f 50%, #d4af37 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin-bottom: 15px;
    }

    .sub-title {
        color: rgba(255, 255, 255, 0.85);
        font-size: 18px;
        letter-spacing: 1px;
    }

    .section-header {
        text-align: center;
        margin: 30px 20px;
        color: #d4af37;
        font-size: 1.3rem;
        font-weight: 700;
    }

    .input-card {
        background: linear-gradient(135deg, rgba(26, 54, 93, 0.9), rgba(15, 30, 50, 0.95));
        border-radius: 16px;
        padding: 25px;
        margin: 10px;
        border: 1px solid rgba(212, 175, 55, 0.2);
    }

    .input-header {
        display: flex;
        align-items: center;
        gap: 15px;
        margin-bottom: 15px;
    }

    .input-icon {
        width: 45px;
        height: 45px;
        display: flex;
        align-items: center;
        justify-content: center;
        background: linear-gradient(135deg, #d4af37, #f4d03f);
        border-radius: 10px;
        font-size: 1.3rem;
    }

    .input-title {
        color: #d4af37;
        font-size: 1.1rem;
        font-weight: 700;
    }

    .input-subtitle {
        color: rgba(255, 255, 255, 0.6);
        font-size: 0.85rem;
    }

    .logo-upload-box {
        background: rgba(255, 255, 255, 0.03);
        border: 2px dashed rgba(212, 175, 55, 0.3);
        border-radius: 12px;
        padding: 20px;
        text-align: center;
        transition: all 0.3s ease;
    }

    .logo-upload-box:hover {
        border-color: #d4af37;
        background: rgba(212, 175, 55, 0.05);
    }

    .logo-preview {
        display: flex;
        align-items: center;
        justify-content: center;
        gap: 10px;
        padding: 15px;
        background: rgba(34, 197, 94, 0.1);
        border-radius: 10px;
        border: 1px solid rgba(34, 197, 94, 0.3);
    }

    .logo-preview img {
        max-height: 60px;
        border-radius: 8px;
    }

    div[role="radiogroup"] {
        display: flex !important;
        flex-direction: row-reverse !important;
        justify-content: center !important;
        gap: 12px !important;
        flex-wrap: wrap !important;
        background: rgba(0, 0, 0, 0.2) !important;
        padding: 20px !important;
        border-radius: 15px !important;
        margin: 0 20px 20px 20px !important;
    }

    div[role="radiogroup"] label {
        background: linear-gradient(135deg, rgba(26, 54, 93, 0.9), rgba(15, 30, 50, 0.95)) !important;
        border: 2px solid rgba(212, 175, 55, 0.2) !important;
        padding: 12px 20px !important;
        border-radius: 10px !important;
        color: white !important;
        font-weight: 600 !important;
        transition: all 0.3s ease !important;
    }

    div[role="radiogroup"] label:hover {
        border-color: #d4af37 !important;
        transform: translateY(-3px) !important;
    }

    .stTextArea textarea {
        background-color: rgba(0, 0, 0, 0.3) !important;
        border: 2px solid rgba(212, 175, 55, 0.2) !important;
        border-radius: 12px !important;
        color: white !important;
        font-family: 'Tajawal', sans-serif !important;
        padding: 15px !important;
        text-align: right !important;
        direction: rtl !important;
    }
    
    .stTextArea textarea:focus {
        border-color: #d4af37 !important;
    }

    .stButton > button {
        background: linear-gradient(135deg, #d4af37, #f4d03f, #d4af37) !important;
        background-size: 200% auto !important;
        color: #1a365d !important;
        font-family: 'Tajawal', sans-serif !important;
        font-weight: 800 !important;
        font-size: 1.2rem !important;
        border-radius: 12px !important;
        width: 100% !important;
        padding: 15px 35px !important;
        border: none !important;
        box-shadow: 0 8px 25px rgba(212, 175, 55, 0.3) !important;
        transition: all 0.3s ease !important;
    }
    
    .stButton > button:hover {
        transform: translateY(-3px) !important;
        box-shadow: 0 12px 35px rgba(212, 175, 55, 0.4) !important;
    }

    .stDownloadButton > button {
        background: linear-gradient(135deg, #1a365d, #2d4a6f) !important;
        color: white !important;
        border: 1px solid rgba(212, 175, 55, 0.3) !important;
    }

    .progress-box {
        background: rgba(26, 54, 93, 0.9);
        border: 1px solid rgba(212, 175, 55, 0.3);
        border-radius: 15px;
        padding: 30px;
        margin: 20px;
        text-align: center;
    }

    .success-banner {
        background: linear-gradient(135deg, rgba(34, 197, 94, 0.2), rgba(34, 197, 94, 0.1));
        border: 2px solid #22c55e;
        border-radius: 12px;
        padding: 20px;
        text-align: center;
        margin: 20px;
    }

    .analysis-card {
        background: linear-gradient(135deg, rgba(26, 54, 93, 0.95), rgba(15, 30, 50, 0.98));
        border: 1px solid rgba(212, 175, 55, 0.3);
        border-radius: 15px;
        padding: 25px;
        margin: 15px 0;
    }
    
    .analysis-title {
        color: #d4af37;
        font-size: 1.2rem;
        font-weight: 700;
        margin-bottom: 20px;
    }

    .stat-item {
        background: rgba(212, 175, 55, 0.1);
        border-radius: 12px;
        padding: 15px;
        text-align: center;
    }
    
    .stat-value {
        font-size: 1.8rem;
        font-weight: 800;
        color: #d4af37;
    }
    
    .stat-label {
        font-size: 0.85rem;
        color: rgba(255, 255, 255, 0.7);
        margin-top: 5px;
    }

    .keyword-tag {
        display: inline-block;
        background: rgba(212, 175, 55, 0.15);
        border: 1px solid rgba(212, 175, 55, 0.3);
        color: #d4af37;
        padding: 5px 12px;
        border-radius: 20px;
        margin: 3px;
        font-size: 0.85rem;
    }

    .stTextArea > label,
    .stFileUploader > label,
    .stRadio > label {
        display: none !important;
    }

    iframe {
        border-radius: 12px !important;
        border: 2px solid rgba(212, 175, 55, 0.3) !important;
    }
</style>
""", unsafe_allow_html=True)

# ---------------------------------------------------------
# 🎨 قوالب CSS للتقارير - تصميم احترافي
# ---------------------------------------------------------

STYLE_OFFICIAL = """
<style>
    @import url('https://fonts.googleapis.com/css2?family=Tajawal:wght@300;400;500;700;800;900&family=Cairo:wght@400;600;700;800&display=swap');
    
    * { box-sizing: border-box; margin: 0; padding: 0; }
    
    body {
        font-family: 'Tajawal', 'Cairo', sans-serif;
        background: #ffffff;
        color: #1a202c;
        line-height: 1.8;
        direction: rtl;
        text-align: right;
    }
    
    .container {
        max-width: 900px;
        margin: 0 auto;
        padding: 40px 50px;
        background: #ffffff;
    }
    
    /* الهيدر مع الشعارين */
    .report-header {
        background: linear-gradient(135deg, #1a365d 0%, #2d4a6f 100%);
        border-radius: 16px;
        padding: 40px;
        margin-bottom: 40px;
        text-align: center;
        position: relative;
        overflow: hidden;
    }
    
    .report-header::before {
        content: '';
        position: absolute;
        top: 0; left: 0; right: 0;
        height: 5px;
        background: linear-gradient(90deg, #d4af37, #f4d03f, #d4af37);
    }
    
    .header-logos {
        display: flex;
        justify-content: space-between;
        align-items: center;
        margin-bottom: 25px;
    }
    
    .header-logos img {
        max-height: 70px;
        filter: brightness(0) invert(1);
    }
    
    .report-header h1 {
        color: #ffffff;
        font-size: 2.2rem;
        font-weight: 800;
        margin-bottom: 10px;
    }
    
    .report-header .subtitle {
        color: #d4af37;
        font-size: 1.1rem;
        font-weight: 500;
    }
    
    .report-header .date {
        color: rgba(255,255,255,0.7);
        font-size: 0.95rem;
        margin-top: 15px;
    }
    
    /* الأقسام */
    .section {
        background: #ffffff;
        border-radius: 12px;
        padding: 30px;
        margin-bottom: 25px;
        border: 1px solid #e2e8f0;
        box-shadow: 0 4px 15px rgba(0,0,0,0.05);
    }
    
    .section-title {
        color: #1a365d;
        font-size: 1.4rem;
        font-weight: 800;
        margin-bottom: 20px;
        padding-bottom: 12px;
        border-bottom: 3px solid #d4af37;
        display: flex;
        align-items: center;
        gap: 10px;
    }
    
    .section-title::before {
        content: '';
        width: 6px;
        height: 25px;
        background: linear-gradient(180deg, #d4af37, #f4d03f);
        border-radius: 3px;
    }
    
    h2 {
        color: #1a365d;
        font-size: 1.3rem;
        font-weight: 700;
        margin: 25px 0 15px 0;
        padding-right: 15px;
        border-right: 4px solid #d4af37;
    }
    
    h3 {
        color: #2d4a6f;
        font-size: 1.1rem;
        font-weight: 700;
        margin: 20px 0 12px 0;
    }
    
    p {
        color: #4a5568;
        font-size: 1.05rem;
        line-height: 1.9;
        margin-bottom: 15px;
    }
    
    /* الجداول */
    table {
        width: 100%;
        border-collapse: separate;
        border-spacing: 0;
        margin: 20px 0;
        border-radius: 10px;
        overflow: hidden;
        box-shadow: 0 2px 10px rgba(0,0,0,0.05);
    }
    
    thead th {
        background: linear-gradient(135deg, #1a365d, #2d4a6f);
        color: #ffffff;
        padding: 15px 18px;
        font-weight: 700;
        text-align: right;
        font-size: 0.95rem;
    }
    
    tbody tr { transition: background 0.2s; }
    tbody tr:nth-child(even) { background: #f8fafc; }
    tbody tr:hover { background: #edf2f7; }
    
    tbody td {
        padding: 14px 18px;
        border-bottom: 1px solid #e2e8f0;
        color: #2d3748;
        font-size: 0.95rem;
    }
    
    /* القوائم */
    ul, ol { list-style: none; padding: 0; margin: 15px 0; }
    
    li {
        padding: 12px 18px;
        margin-bottom: 8px;
        background: linear-gradient(135deg, #f8fafc, #ffffff);
        border-radius: 8px;
        border-right: 4px solid #d4af37;
        color: #2d3748;
        font-size: 1rem;
        transition: all 0.2s;
    }
    
    li:hover {
        transform: translateX(-5px);
        box-shadow: 0 4px 12px rgba(0,0,0,0.08);
    }
    
    /* بطاقات الإحصائيات */
    .stats-grid {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(160px, 1fr));
        gap: 15px;
        margin: 25px 0;
    }
    
    .stat-card {
        background: linear-gradient(135deg, #f8fafc, #ffffff);
        border-radius: 12px;
        padding: 25px 20px;
        text-align: center;
        border: 1px solid #e2e8f0;
        transition: all 0.3s;
    }
    
    .stat-card:hover {
        transform: translateY(-5px);
        box-shadow: 0 10px 25px rgba(0,0,0,0.1);
    }
    
    .stat-card .value {
        font-size: 2.2rem;
        font-weight: 900;
        color: #1a365d;
        display: block;
    }
    
    .stat-card .label {
        font-size: 0.9rem;
        color: #718096;
        margin-top: 5px;
    }
    
    /* صندوق التمييز */
    .highlight-box {
        background: linear-gradient(135deg, #fef9e7, #fdf6e3);
        border-radius: 12px;
        padding: 20px 25px;
        margin: 20px 0;
        border-right: 5px solid #d4af37;
    }
    
    .highlight-box.info {
        background: linear-gradient(135deg, #ebf8ff, #e6f7ff);
        border-right-color: #3182ce;
    }
    
    .highlight-box.success {
        background: linear-gradient(135deg, #f0fff4, #e6ffed);
        border-right-color: #38a169;
    }
    
    .highlight-box.warning {
        background: linear-gradient(135deg, #fffaf0, #ffefd5);
        border-right-color: #dd6b20;
    }
    
    /* الفوتر */
    footer {
        text-align: center;
        margin-top: 40px;
        padding: 30px;
        background: linear-gradient(135deg, #f8fafc, #ffffff);
        border-radius: 12px;
        border-top: 4px solid #d4af37;
    }
    
    footer .org-name {
        color: #1a365d;
        font-weight: 800;
        font-size: 1.1rem;
    }
    
    footer .dept-name {
        color: #d4af37;
        font-size: 1rem;
        margin-top: 5px;
    }
    
    footer .copyright {
        color: #a0aec0;
        font-size: 0.85rem;
        margin-top: 15px;
    }
    
    /* طباعة */
    @media print {
        body { background: white; }
        .section { box-shadow: none; page-break-inside: avoid; }
        .report-header { -webkit-print-color-adjust: exact; print-color-adjust: exact; }
    }
</style>
"""

STYLE_DIGITAL = """
<style>
    @import url('https://fonts.googleapis.com/css2?family=Tajawal:wght@300;400;500;700;800;900&family=Cairo:wght@400;600;700;800&display=swap');
    
    * { box-sizing: border-box; margin: 0; padding: 0; }
    
    body {
        font-family: 'Tajawal', 'Cairo', sans-serif;
        background: #ffffff;
        color: #1a1a2e;
        line-height: 1.8;
        direction: rtl;
    }
    
    .container {
        max-width: 950px;
        margin: 0 auto;
        padding: 40px 50px;
        background: #ffffff;
    }
    
    /* الهيدر */
    .report-header {
        background: linear-gradient(135deg, #0066cc 0%, #004999 100%);
        border-radius: 20px;
        padding: 45px;
        margin-bottom: 40px;
        text-align: center;
        position: relative;
    }
    
    .report-header::after {
        content: '';
        position: absolute;
        bottom: -15px;
        left: 50%;
        transform: translateX(-50%);
        width: 80px;
        height: 5px;
        background: #00c853;
        border-radius: 3px;
    }
    
    .header-logos {
        display: flex;
        justify-content: space-between;
        align-items: center;
        margin-bottom: 25px;
    }
    
    .header-logos img {
        max-height: 65px;
        filter: brightness(0) invert(1);
    }
    
    .report-header h1 {
        color: #ffffff;
        font-size: 2.3rem;
        font-weight: 800;
        margin-bottom: 12px;
    }
    
    .report-header .subtitle {
        color: rgba(255,255,255,0.9);
        font-size: 1.1rem;
    }
    
    /* البطاقات */
    .card {
        background: #ffffff;
        border-radius: 16px;
        padding: 30px;
        margin-bottom: 25px;
        border: 1px solid #e8ecf0;
        box-shadow: 0 4px 20px rgba(0,0,0,0.04);
    }
    
    .card-title {
        color: #0066cc;
        font-size: 1.35rem;
        font-weight: 800;
        margin-bottom: 20px;
        display: flex;
        align-items: center;
        gap: 12px;
    }
    
    .card-title .icon {
        width: 40px;
        height: 40px;
        background: #e6f2ff;
        border-radius: 10px;
        display: flex;
        align-items: center;
        justify-content: center;
    }
    
    h2 {
        color: #0066cc;
        font-size: 1.4rem;
        font-weight: 800;
        margin: 30px 0 18px 0;
        padding-bottom: 10px;
        border-bottom: 2px solid #e8ecf0;
    }
    
    h3 {
        color: #1a1a2e;
        font-size: 1.15rem;
        font-weight: 700;
        margin: 20px 0 12px 0;
    }
    
    p {
        color: #4a4a68;
        font-size: 1.05rem;
        line-height: 1.9;
    }
    
    /* شبكة الإحصائيات */
    .metrics-grid {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(170px, 1fr));
        gap: 18px;
        margin: 25px 0;
    }
    
    .metric-card {
        background: #f8fafc;
        border-radius: 14px;
        padding: 25px;
        text-align: center;
        border: 1px solid #e8ecf0;
        position: relative;
        overflow: hidden;
    }
    
    .metric-card::before {
        content: '';
        position: absolute;
        top: 0; left: 0; right: 0;
        height: 4px;
        background: #0066cc;
    }
    
    .metric-card.success::before { background: #00c853; }
    .metric-card.warning::before { background: #ff9100; }
    .metric-card.danger::before { background: #f44336; }
    
    .metric-card .number {
        font-size: 2.4rem;
        font-weight: 900;
        color: #0066cc;
        display: block;
    }
    
    .metric-card .label {
        font-size: 0.9rem;
        color: #8888a0;
        margin-top: 5px;
    }
    
    /* الجداول */
    table {
        width: 100%;
        border-collapse: separate;
        border-spacing: 0;
        margin: 25px 0;
        border-radius: 12px;
        overflow: hidden;
    }
    
    thead th {
        background: #0066cc;
        color: white;
        padding: 15px 18px;
        font-weight: 700;
        text-align: right;
    }
    
    tbody tr:nth-child(even) { background: #f8fafc; }
    tbody tr:hover { background: #e6f2ff; }
    
    tbody td {
        padding: 14px 18px;
        border-bottom: 1px solid #e8ecf0;
    }
    
    /* القوائم */
    ul { list-style: none; padding: 0; margin: 18px 0; }
    
    ul li {
        padding: 14px 20px;
        margin-bottom: 10px;
        background: #f8fafc;
        border-radius: 10px;
        position: relative;
        padding-right: 45px;
        transition: all 0.2s;
    }
    
    ul li::before {
        content: '✓';
        position: absolute;
        right: 15px;
        top: 50%;
        transform: translateY(-50%);
        color: #00c853;
        font-weight: bold;
        font-size: 1.1rem;
    }
    
    ul li:hover { background: #e6f2ff; }
    
    /* صندوق التمييز */
    .highlight-box {
        background: linear-gradient(135deg, #e6f2ff, #f0f7ff);
        border-radius: 12px;
        padding: 22px;
        margin: 22px 0;
        border-right: 5px solid #0066cc;
    }
    
    .highlight-box.success {
        background: linear-gradient(135deg, #e8f5e9, #f1f8e9);
        border-right-color: #00c853;
    }
    
    /* الفوتر */
    footer {
        text-align: center;
        margin-top: 40px;
        padding: 30px;
        background: #f8fafc;
        border-radius: 16px;
    }
    
    footer .org-name {
        color: #0066cc;
        font-weight: 800;
        font-size: 1.1rem;
    }
    
    footer .dept-name {
        color: #4a4a68;
        font-size: 1rem;
        margin-top: 5px;
    }
</style>
"""

STYLE_ANALYTICAL = """
<style>
    @import url('https://fonts.googleapis.com/css2?family=Tajawal:wght@300;400;500;700;800;900&family=Cairo:wght@400;600;700;800&display=swap');
    
    * { box-sizing: border-box; margin: 0; padding: 0; }
    
    body {
        font-family: 'Tajawal', 'Cairo', sans-serif;
        background: #ffffff;
        color: #1a1a2e;
        line-height: 1.8;
        direction: rtl;
    }
    
    .container {
        max-width: 950px;
        margin: 0 auto;
        padding: 40px 50px;
        background: #ffffff;
    }
    
    /* الهيدر */
    .report-header {
        background: linear-gradient(135deg, #1e3a5f 0%, #3d5a80 100%);
        border-radius: 18px;
        padding: 45px;
        margin-bottom: 40px;
        text-align: center;
    }
    
    .header-logos {
        display: flex;
        justify-content: space-between;
        align-items: center;
        margin-bottom: 25px;
    }
    
    .header-logos img {
        max-height: 65px;
        filter: brightness(0) invert(1);
    }
    
    .report-header h1 {
        color: #ffffff;
        font-size: 2.2rem;
        font-weight: 800;
        margin-bottom: 12px;
    }
    
    .report-header .subtitle {
        color: #ee6c4d;
        font-size: 1.1rem;
        font-weight: 500;
    }
    
    /* أقسام التقرير */
    .report-section {
        background: #ffffff;
        border-radius: 14px;
        padding: 30px;
        margin-bottom: 25px;
        border: 1px solid #e9ecef;
        box-shadow: 0 4px 15px rgba(0,0,0,0.04);
    }
    
    .section-header {
        display: flex;
        align-items: center;
        gap: 15px;
        margin-bottom: 25px;
        padding-bottom: 15px;
        border-bottom: 2px solid #e9ecef;
    }
    
    .section-number {
        width: 45px;
        height: 45px;
        background: linear-gradient(135deg, #ee6c4d, #ff8a65);
        border-radius: 10px;
        display: flex;
        align-items: center;
        justify-content: center;
        color: white;
        font-size: 1.3rem;
        font-weight: 800;
    }
    
    .section-header h2 {
        color: #1e3a5f;
        font-size: 1.4rem;
        font-weight: 800;
        margin: 0;
    }
    
    h3 {
        color: #3d5a80;
        font-size: 1.15rem;
        font-weight: 700;
        margin: 20px 0 12px 0;
    }
    
    p {
        color: #4a5568;
        font-size: 1.05rem;
        line-height: 1.9;
    }
    
    /* شبكة الإحصائيات */
    .stats-grid {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(180px, 1fr));
        gap: 18px;
        margin: 25px 0;
    }
    
    .stat-card {
        background: #f8f9fa;
        border-radius: 14px;
        padding: 28px;
        text-align: center;
        position: relative;
        overflow: hidden;
        border: 1px solid #e9ecef;
    }
    
    .stat-card::after {
        content: '';
        position: absolute;
        bottom: 0; left: 0; right: 0;
        height: 4px;
        background: #1e3a5f;
    }
    
    .stat-card.highlight::after { background: #ee6c4d; }
    .stat-card.success::after { background: #06d6a0; }
    .stat-card.warning::after { background: #ffd166; }
    
    .stat-card .value {
        font-size: 2.5rem;
        font-weight: 900;
        color: #1e3a5f;
        display: block;
    }
    
    .stat-card .label {
        font-size: 0.9rem;
        color: #718096;
        margin-top: 5px;
    }
    
    /* شريط التقدم */
    .progress-item { margin-bottom: 18px; }
    
    .progress-header {
        display: flex;
        justify-content: space-between;
        margin-bottom: 8px;
    }
    
    .progress-label { color: #1a1a2e; font-weight: 600; }
    .progress-value { color: #1e3a5f; font-weight: 800; }
    
    .progress-bar {
        height: 10px;
        background: #e9ecef;
        border-radius: 5px;
        overflow: hidden;
    }
    
    .progress-fill {
        height: 100%;
        background: linear-gradient(90deg, #1e3a5f, #3d5a80);
        border-radius: 5px;
    }
    
    .progress-fill.success { background: linear-gradient(90deg, #06d6a0, #00e5b0); }
    .progress-fill.warning { background: linear-gradient(90deg, #ffd166, #ffe066); }
    .progress-fill.danger { background: linear-gradient(90deg, #ee6c4d, #ff8a80); }
    
    /* الجداول */
    table {
        width: 100%;
        border-collapse: separate;
        border-spacing: 0;
        margin: 25px 0;
        border-radius: 12px;
        overflow: hidden;
    }
    
    thead th {
        background: #1e3a5f;
        color: white;
        padding: 15px 18px;
        font-weight: 700;
        text-align: right;
    }
    
    tbody tr:nth-child(even) { background: #f8f9fa; }
    tbody tr:hover { background: #eef2f7; }
    
    tbody td {
        padding: 14px 18px;
        border-bottom: 1px solid #e9ecef;
    }
    
    /* البطاقات المميزة */
    .highlight-card {
        background: #ffeae5;
        border-radius: 12px;
        padding: 22px;
        margin: 22px 0;
        border-right: 5px solid #ee6c4d;
    }
    
    .highlight-card.info {
        background: #e3f2fd;
        border-right-color: #3d5a80;
    }
    
    .highlight-card.success {
        background: #e8f5e9;
        border-right-color: #06d6a0;
    }
    
    /* الفوتر */
    footer {
        text-align: center;
        margin-top: 40px;
        padding: 30px;
        background: #f8f9fa;
        border-radius: 14px;
        border-top: 4px solid #ee6c4d;
    }
    
    footer .org-name {
        color: #1e3a5f;
        font-weight: 800;
        font-size: 1.1rem;
    }
    
    footer .dept-name {
        color: #ee6c4d;
        font-size: 1rem;
        margin-top: 5px;
    }
</style>
"""

STYLE_PRESENTATION = """
<style>
    @import url('https://fonts.googleapis.com/css2?family=Tajawal:wght@300;400;500;700;800;900&family=Cairo:wght@400;600;700;800&display=swap');
    
    * { box-sizing: border-box; margin: 0; padding: 0; }
    
    body {
        font-family: 'Tajawal', 'Cairo', sans-serif;
        background: #ffffff;
        overflow: hidden;
        height: 100vh;
        width: 100vw;
        direction: rtl;
    }
    
    .presentation-container {
        width: 100%;
        height: 100%;
        position: relative;
        background: #ffffff;
    }
    
    .slide {
        position: absolute;
        top: 0; left: 0;
        width: 100%;
        height: 100%;
        opacity: 0;
        visibility: hidden;
        transform: scale(0.95);
        transition: all 0.5s ease;
        display: flex;
        flex-direction: column;
        padding: 50px 60px;
        background: #ffffff;
    }
    
    .slide.active {
        opacity: 1;
        visibility: visible;
        transform: scale(1);
        z-index: 10;
    }
    
    /* شريحة الغلاف */
    .slide.cover {
        align-items: center;
        justify-content: center;
        text-align: center;
        background: linear-gradient(135deg, #1a365d 0%, #2d4a6f 100%);
    }
    
    .cover-content {
        padding: 70px;
        border: 3px solid #d4af37;
        border-radius: 20px;
        background: rgba(255,255,255,0.05);
    }
    
    .cover-logos {
        display: flex;
        justify-content: center;
        align-items: center;
        gap: 40px;
        margin-bottom: 30px;
    }
    
    .cover-logos img { max-height: 80px; }
    
    .slide.cover .main-title {
        font-size: 3rem;
        color: white;
        font-weight: 900;
        margin-bottom: 20px;
    }
    
    .slide.cover .sub-title {
        font-size: 1.4rem;
        color: #d4af37;
    }
    
    /* الشرائح العادية */
    .slide-header {
        display: flex;
        justify-content: space-between;
        align-items: center;
        padding-bottom: 20px;
        margin-bottom: 30px;
        border-bottom: 3px solid #d4af37;
    }
    
    .header-title h2 {
        color: #1a365d;
        font-size: 1.8rem;
        font-weight: 800;
    }
    
    .header-logos {
        display: flex;
        gap: 20px;
    }
    
    .header-logos img { max-height: 45px; }
    
    .slide-content {
        flex-grow: 1;
        display: flex;
        gap: 40px;
        overflow: hidden;
    }
    
    .text-panel {
        flex: 3;
        padding: 25px;
        overflow-y: auto;
    }
    
    .visual-panel {
        flex: 2;
        display: flex;
        flex-direction: column;
        justify-content: center;
        align-items: center;
        background: linear-gradient(135deg, #1a365d, #2d4a6f);
        border-radius: 20px;
        color: white;
        padding: 35px;
    }
    
    h3 {
        color: #1a365d;
        font-size: 1.4rem;
        margin-bottom: 18px;
        font-weight: 700;
    }
    
    p {
        font-size: 1.15rem;
        line-height: 1.9;
        color: #2d3748;
        margin-bottom: 18px;
    }
    
    li {
        font-size: 1.1rem;
        margin-bottom: 12px;
        color: #2d3748;
    }
    
    .icon-box {
        font-size: 4.5rem;
        color: #d4af37;
        margin-bottom: 20px;
    }
    
    /* التنقل */
    .nav-controls {
        position: absolute;
        bottom: 25px;
        left: 50%;
        transform: translateX(-50%);
        display: flex;
        gap: 15px;
        z-index: 100;
    }
    
    .nav-btn {
        width: 48px;
        height: 48px;
        border-radius: 50%;
        border: 2px solid #1a365d;
        background: white;
        color: #1a365d;
        cursor: pointer;
        font-size: 1.1rem;
        transition: all 0.3s;
    }
    
    .nav-btn:hover {
        background: #1a365d;
        color: white;
    }
    
    .page-number {
        position: absolute;
        bottom: 30px;
        right: 60px;
        color: #718096;
        font-size: 1rem;
        font-weight: 600;
    }
</style>
<link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css">
"""

STYLE_EXECUTIVE = """
<style>
    @import url('https://fonts.googleapis.com/css2?family=Tajawal:wght@300;400;500;700;800;900&family=Cairo:wght@400;600;700;800&display=swap');
    
    * { box-sizing: border-box; margin: 0; padding: 0; }
    
    body {
        font-family: 'Tajawal', 'Cairo', sans-serif;
        background: #ffffff;
        color: #111827;
        line-height: 1.8;
        direction: rtl;
    }
    
    .container {
        max-width: 850px;
        margin: 0 auto;
        padding: 50px 60px;
        background: #ffffff;
    }
    
    /* الهيدر */
    header {
        display: flex;
        justify-content: space-between;
        align-items: center;
        padding-bottom: 25px;
        margin-bottom: 40px;
        border-bottom: 4px solid #111827;
    }
    
    .header-logos {
        display: flex;
        gap: 20px;
    }
    
    .header-logos img { max-height: 55px; }
    
    .brand {
        font-size: 1.2rem;
        font-weight: 800;
        color: #111827;
    }
    
    /* العنوان الرئيسي */
    h1 {
        font-size: 2.8rem;
        font-weight: 900;
        color: #111827;
        line-height: 1.2;
        margin-bottom: 15px;
    }
    
    .subtitle {
        font-size: 1.15rem;
        color: #6b7280;
        margin-bottom: 35px;
    }
    
    /* الملخص التنفيذي */
    .executive-summary {
        font-size: 1.2rem;
        line-height: 1.9;
        color: #374151;
        margin-bottom: 45px;
        padding: 28px;
        background: #f9fafb;
        border-radius: 12px;
        border-right: 5px solid #d97706;
    }
    
    /* شبكة المقاييس */
    .metrics-grid {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(170px, 1fr));
        gap: 22px;
        margin: 35px 0;
    }
    
    .metric-box {
        padding: 28px;
        background: #f9fafb;
        border-radius: 12px;
        text-align: center;
        transition: all 0.3s;
    }
    
    .metric-box:hover {
        transform: translateY(-4px);
        box-shadow: 0 10px 25px rgba(0,0,0,0.08);
    }
    
    .metric-value {
        font-size: 2.6rem;
        font-weight: 900;
        color: #111827;
        display: block;
    }
    
    .metric-label {
        font-size: 0.85rem;
        color: #6b7280;
        text-transform: uppercase;
        letter-spacing: 1px;
        margin-top: 5px;
    }
    
    /* عناوين الأقسام */
    .section-title {
        font-size: 1rem;
        font-weight: 800;
        text-transform: uppercase;
        letter-spacing: 2px;
        color: #d97706;
        margin: 45px 0 22px 0;
        padding-bottom: 12px;
        border-bottom: 2px solid #e5e7eb;
    }
    
    h2 {
        font-size: 1.7rem;
        font-weight: 800;
        color: #111827;
        margin: 30px 0 18px 0;
    }
    
    h3 {
        font-size: 1.25rem;
        font-weight: 700;
        color: #374151;
        margin: 22px 0 14px 0;
    }
    
    p {
        font-size: 1.1rem;
        color: #4b5563;
        line-height: 1.9;
        margin-bottom: 18px;
    }
    
    /* القوائم */
    ul, ol { list-style: none; padding: 0; margin: 18px 0; }
    
    li {
        padding: 12px 0;
        border-bottom: 1px solid #e5e7eb;
        display: flex;
        align-items: center;
        gap: 12px;
        color: #374151;
    }
    
    li::before {
        content: '—';
        color: #d97706;
        font-weight: 800;
    }
    
    /* الجداول */
    table {
        width: 100%;
        border-collapse: collapse;
        margin: 28px 0;
    }
    
    thead th {
        background: #111827;
        color: white;
        padding: 15px 18px;
        text-align: right;
        font-weight: 700;
    }
    
    tbody tr:nth-child(even) { background: #f9fafb; }
    
    tbody td {
        padding: 14px 18px;
        border-bottom: 1px solid #e5e7eb;
    }
    
    /* صندوق التمييز */
    .callout {
        background: #fef3c7;
        padding: 22px 28px;
        border-radius: 12px;
        margin: 28px 0;
        border-right: 5px solid #d97706;
    }
    
    .callout.success {
        background: #d1fae5;
        border-right-color: #059669;
    }
    
    /* الفوتر */
    footer {
        margin-top: 60px;
        padding-top: 28px;
        border-top: 2px solid #e5e7eb;
        text-align: center;
    }
    
    footer .org-name {
        font-size: 1.05rem;
        font-weight: 800;
        color: #111827;
    }
    
    footer .dept-name {
        font-size: 0.95rem;
        color: #d97706;
        margin-top: 5px;
    }
    
    footer .date {
        font-size: 0.85rem;
        color: #9ca3af;
        margin-top: 15px;
    }
</style>
"""

SCRIPT_PRESENTATION = """
<script>
    let currentSlideIndex = 1;
    function updateSlide() {
        const slides = document.querySelectorAll('.slide');
        const totalSlides = slides.length;
        if(totalSlides === 0) return;
        slides.forEach(slide => { slide.classList.remove('active'); });
        const activeSlide = document.getElementById(`slide-${currentSlideIndex}`);
        if(activeSlide) activeSlide.classList.add('active');
        const pageNum = document.getElementById('page-num');
        if(pageNum) pageNum.innerText = `${currentSlideIndex} / ${totalSlides}`;
    }
    function nextSlide() {
        const totalSlides = document.querySelectorAll('.slide').length;
        if (currentSlideIndex < totalSlides) { currentSlideIndex++; updateSlide(); }
    }
    function prevSlide() {
        if (currentSlideIndex > 1) { currentSlideIndex--; updateSlide(); }
    }
    document.addEventListener('keydown', function(event) {
        if (event.key === "ArrowLeft" || event.key === " ") nextSlide();
        else if (event.key === "ArrowRight") prevSlide();
    });
    setTimeout(updateSlide, 100);
</script>
"""

# ---------------------------------------------------------
# 🛠️ دوال المساعدة
# ---------------------------------------------------------

def extract_text_from_file(uploaded_file):
    """استخراج النص من الملفات"""
    text_content = ""
    try:
        if uploaded_file.type == "application/pdf":
            reader = PyPDF2.PdfReader(uploaded_file)
            for page in reader.pages:
                text_content += page.extract_text() + "\n"
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            df = pd.read_excel(uploaded_file)
            text_content = df.to_string()
        else:
            stringio = StringIO(uploaded_file.getvalue().decode("utf-8"))
            text_content = stringio.read()
    except Exception as e:
        return f"خطأ في قراءة الملف: {e}"
    return text_content

def clean_html_response(text):
    """تنظيف رد الذكاء الاصطناعي"""
    text = text.replace("```html", "").replace("```", "")
    return text.strip()

def get_working_model():
    """الحصول على نموذج Gemini المتاح"""
    try:
        available_models = []
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                available_models.append(m.name)
        
        preferred = ['models/gemini-1.5-flash-latest', 'models/gemini-1.5-flash', 'models/gemini-1.5-pro-latest', 'models/gemini-1.5-pro', 'models/gemini-pro']
        
        for model in preferred:
            if model in available_models:
                return model
        
        if available_models:
            return available_models[0]
        return 'gemini-pro'
    except:
        return 'gemini-pro'

def analyze_text(text):
    """تحليل النص واستخراج الإحصائيات"""
    if not text.strip():
        return None
    
    clean_text = re.sub(r'[^\w\s]', '', text)
    words = clean_text.split()
    word_count = len(words)
    
    sentences = re.split(r'[.!?؟،]', text)
    sentence_count = len([s for s in sentences if s.strip()])
    
    paragraphs = text.split('\n\n')
    paragraph_count = len([p for p in paragraphs if p.strip()])
    
    char_count = len(text)
    numbers = re.findall(r'\d+(?:\.\d+)?%?', text)
    
    stop_words = {'في', 'من', 'إلى', 'على', 'عن', 'مع', 'هذا', 'هذه', 'التي', 'الذي', 'أن', 'كان', 'كانت', 'يكون', 'تكون', 'هو', 'هي', 'ذلك', 'تلك', 'و', 'أو', 'ثم', 'لكن', 'بل', 'حتى', 'إذا', 'لو', 'ما', 'لا', 'نعم', 'قد', 'لقد', 'سوف', 'عند', 'بعد', 'قبل'}
    
    filtered_words = [w for w in words if len(w) > 2 and w not in stop_words]
    word_freq = Counter(filtered_words)
    keywords = word_freq.most_common(8)
    
    return {
        'word_count': word_count,
        'sentence_count': sentence_count,
        'paragraph_count': paragraph_count,
        'char_count': char_count,
        'numbers': numbers[:10],
        'keywords': keywords,
        'reading_time': max(1, word_count // 200)
    }

def image_to_base64(uploaded_image):
    """تحويل الصورة إلى Base64"""
    if uploaded_image is not None:
        bytes_data = uploaded_image.getvalue()
        base64_str = base64.b64encode(bytes_data).decode()
        return f"data:image/{uploaded_image.type.split('/')[-1]};base64,{base64_str}"
    return None

def create_pdf(html_content):
    """إنشاء ملف PDF من HTML"""
    if not WEASYPRINT_AVAILABLE:
        return None
    try:
        pdf_bytes = WeasyHTML(string=html_content).write_pdf()
        return pdf_bytes
    except Exception as e:
        st.error(f"خطأ في إنشاء PDF: {e}")
        return None

def create_docx(html_content, title="تقرير"):
    """إنشاء ملف Word من HTML"""
    if not DOCX_AVAILABLE:
        return None
    try:
        doc = Document()
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        text_content = re.sub(r'<[^>]+>', '\n', html_content)
        text_content = re.sub(r'\n\s*\n', '\n\n', text_content)
        
        for para in text_content.split('\n\n'):
            if para.strip():
                p = doc.add_paragraph(para.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    except Exception as e:
        st.error(f"خطأ في إنشاء Word: {e}")
        return None

# ---------------------------------------------------------
# 🏗️ بناء الواجهة
# ---------------------------------------------------------

# الهيدر
st.markdown('''
<div class="hero-section">
    <div class="main-title">تيار الحكمة الوطني</div>
    <div class="sub-title">الجهاز المركزي للجودة الشاملة | وحدة التخطيط الاستراتيجي</div>
</div>
''', unsafe_allow_html=True)

# ===== قسم الشعارات (شعارين) =====
st.markdown('<div class="section-header">🏷️ شعارات المؤسسة (اختياري - يمكنك رفع شعارين)</div>', unsafe_allow_html=True)

logo_col1, logo_col2 = st.columns(2)

with logo_col1:
    st.markdown('''
    <div class="input-card">
        <div class="input-header">
            <div class="input-icon">🏛️</div>
            <div>
                <div class="input-title">الشعار الأيمن</div>
                <div class="input-subtitle">شعار الجهة الرئيسية</div>
            </div>
        </div>
    </div>
    ''', unsafe_allow_html=True)
    uploaded_logo_right = st.file_uploader("شعار أيمن", type=['png', 'jpg', 'jpeg'], label_visibility="collapsed", key="logo_right")
    
    if uploaded_logo_right:
        logo_right_b64 = image_to_base64(uploaded_logo_right)
        st.markdown(f'<div class="logo-preview"><img src="{logo_right_b64}" alt="شعار"><span style="color: #22c55e;">✅</span></div>', unsafe_allow_html=True)

with logo_col2:
    st.markdown('''
    <div class="input-card">
        <div class="input-header">
            <div class="input-icon">🎯</div>
            <div>
                <div class="input-title">الشعار الأيسر</div>
                <div class="input-subtitle">شعار إضافي (اختياري)</div>
            </div>
        </div>
    </div>
    ''', unsafe_allow_html=True)
    uploaded_logo_left = st.file_uploader("شعار أيسر", type=['png', 'jpg', 'jpeg'], label_visibility="collapsed", key="logo_left")
    
    if uploaded_logo_left:
        logo_left_b64 = image_to_base64(uploaded_logo_left)
        st.markdown(f'<div class="logo-preview"><img src="{logo_left_b64}" alt="شعار"><span style="color: #22c55e;">✅</span></div>', unsafe_allow_html=True)

st.markdown("<br>", unsafe_allow_html=True)

# ===== قسم اختيار النمط =====
st.markdown('<div class="section-header">🎨 اختر نمط الإخراج المطلوب</div>', unsafe_allow_html=True)

report_type = st.radio(
    "",
    ("🏛️ نمط الكتاب الرسمي", "📱 نمط الداشبورد الرقمي", "📊 نمط التحليل العميق", "📽️ عرض تقديمي تفاعلي", "✨ ملخص تنفيذي حديث"),
    horizontal=True,
    label_visibility="collapsed"
)

st.markdown("<br>", unsafe_allow_html=True)

# ===== قسم الإدخال =====
col_input, col_upload = st.columns([2, 1])

with col_input:
    st.markdown('''
    <div class="input-card">
        <div class="input-header">
            <div class="input-icon">📝</div>
            <div>
                <div class="input-title">البيانات / الملاحظات</div>
                <div class="input-subtitle">أدخل النص أو الصق محتوى التقرير هنا</div>
            </div>
        </div>
    </div>
    ''', unsafe_allow_html=True)
    
    user_text = st.text_area("", height=250, placeholder="اكتب الملاحظات أو الصق نص التقرير هنا...", label_visibility="collapsed")

with col_upload:
    st.markdown('''
    <div class="input-card">
        <div class="input-header">
            <div class="input-icon">📎</div>
            <div>
                <div class="input-title">رفع الملفات</div>
                <div class="input-subtitle">PDF, XLSX, TXT</div>
            </div>
        </div>
    </div>
    ''', unsafe_allow_html=True)
    uploaded_file = st.file_uploader("", type=['pdf', 'xlsx', 'txt'], label_visibility="collapsed", key="file_uploader")
    
    if uploaded_file:
        st.success(f"✅ تم إرفاق: {uploaded_file.name}")

# ===== تحليل النص الذكي =====
if user_text.strip():
    analysis = analyze_text(user_text)
    if analysis:
        st.markdown('<div class="analysis-card"><div class="analysis-title">📊 تحليل النص الذكي</div></div>', unsafe_allow_html=True)
        
        stat_cols = st.columns(5)
        stats = [
            (analysis['word_count'], 'كلمة'),
            (analysis['sentence_count'], 'جملة'),
            (analysis['paragraph_count'], 'فقرة'),
            (analysis['char_count'], 'حرف'),
            (analysis['reading_time'], 'دقائق للقراءة')
        ]
        
        for i, (val, label) in enumerate(stats):
            with stat_cols[i]:
                st.markdown(f'<div class="stat-item"><div class="stat-value">{val}</div><div class="stat-label">{label}</div></div>', unsafe_allow_html=True)
        
        if analysis['keywords']:
            keywords_html = " ".join([f'<span class="keyword-tag">{word} ({count})</span>' for word, count in analysis['keywords']])
            st.markdown(f'<p style="color: #d4af37; margin: 15px 0 10px 0; font-weight: 600;">🔑 الكلمات المفتاحية:</p><div>{keywords_html}</div>', unsafe_allow_html=True)

st.markdown("<br>", unsafe_allow_html=True)

# ===== زر المعالجة =====
if st.button("🚀 بدء المعالجة وإنشاء التقرير الكامل"):
    
    if not API_KEY:
        st.error("⚠️ لم يتم العثور على مفتاح API. يرجى إضافته في Secrets.")
        st.stop()
    
    full_text = user_text
    if uploaded_file:
        with st.spinner('📂 جاري قراءة الملف...'):
            full_text += f"\n\n[محتوى الملف]:\n{extract_text_from_file(uploaded_file)}"

    if not full_text.strip():
        st.warning("⚠️ الرجاء إدخال بيانات أو رفع ملف.")
    else:
        try:
            genai.configure(api_key=API_KEY)
            model_name = get_working_model()
            model = genai.GenerativeModel(model_name)

            # إعداد الشعارات
            logo_right_html = ""
            logo_left_html = ""
            
            if uploaded_logo_right:
                logo_right_b64 = image_to_base64(uploaded_logo_right)
                logo_right_html = f'<img src="{logo_right_b64}" alt="شعار" style="max-height: 70px;">'
            
            if uploaded_logo_left:
                logo_left_b64 = image_to_base64(uploaded_logo_left)
                logo_left_html = f'<img src="{logo_left_b64}" alt="شعار" style="max-height: 70px;">'
            
            # إنشاء HTML للشعارات
            logos_html = ""
            if logo_right_html or logo_left_html:
                logos_html = f'''
                <div class="header-logos">
                    <div>{logo_right_html}</div>
                    <div>{logo_left_html}</div>
                </div>
                '''

            # تحديد النمط والقواعد
            target_css = ""
            design_rules = ""
            file_label = "Report"

            if "الرسمي" in report_type:
                target_css = STYLE_OFFICIAL
                file_label = "Official_Report"
                design_rules = f"""
                Create a professional official report with this structure:
                
                1. START with this exact header structure:
                <div class="report-header">
                    {logos_html}
                    <h1>[عنوان التقرير]</h1>
                    <p class="subtitle">[العنوان الفرعي]</p>
                    <p class="date">[التاريخ]</p>
                </div>
                
                2. Wrap each major section in: <div class="section"><h2 class="section-title">[عنوان]</h2>...</div>
                
                3. Use these elements:
                   - <div class="stats-grid"><div class="stat-card"><span class="value">XX</span><span class="label">وصف</span></div></div> for statistics
                   - <table><thead><tr><th>...</th></tr></thead><tbody><tr><td>...</td></tr></tbody></table> for tables
                   - <ul><li>...</li></ul> for lists
                   - <div class="highlight-box">...</div> for important notes
                
                4. END with:
                <footer>
                    <p class="org-name">الجهاز المركزي للجودة الشاملة</p>
                    <p class="dept-name">وحدة التخطيط الاستراتيجي والتطوير</p>
                </footer>
                """
            
            elif "الرقمي" in report_type:
                target_css = STYLE_DIGITAL
                file_label = "Digital_Dashboard"
                design_rules = f"""
                Create a modern digital dashboard report:
                
                1. START with:
                <div class="report-header">
                    {logos_html}
                    <h1>[عنوان التقرير]</h1>
                    <p class="subtitle">[الوصف]</p>
                </div>
                
                2. Use <div class="card"><h2 class="card-title">[عنوان]</h2>...</div> for sections
                
                3. Use these elements:
                   - <div class="metrics-grid"><div class="metric-card"><span class="number">XX</span><span class="label">وصف</span></div></div>
                   - Add class "success", "warning", or "danger" to metric-card for colors
                   - <div class="highlight-box">...</div> or <div class="highlight-box success">...</div>
                
                4. END with footer
                """
            
            elif "التحليل" in report_type:
                target_css = STYLE_ANALYTICAL
                file_label = "Deep_Analysis"
                design_rules = f"""
                Create an analytical deep-dive report:
                
                1. START with:
                <div class="report-header">
                    {logos_html}
                    <h1>[عنوان التقرير]</h1>
                    <p class="subtitle">[الوصف]</p>
                </div>
                
                2. Number each section:
                <div class="report-section">
                    <div class="section-header">
                        <div class="section-number">1</div>
                        <h2>[عنوان القسم]</h2>
                    </div>
                    [المحتوى]
                </div>
                
                3. Use:
                   - <div class="stats-grid"><div class="stat-card highlight"><span class="value">XX%</span><span class="label">وصف</span></div></div>
                   - Progress bars: <div class="progress-item"><div class="progress-header"><span class="progress-label">اسم</span><span class="progress-value">75%</span></div><div class="progress-bar"><div class="progress-fill" style="width: 75%;"></div></div></div>
                   - <div class="highlight-card">...</div>
                
                4. END with footer
                """
            
            elif "ملخص" in report_type:
                target_css = STYLE_EXECUTIVE
                file_label = "Executive_Summary"
                design_rules = f"""
                Create a minimalist executive summary:
                
                1. START with:
                <header>
                    <div class="header-logos">
                        {logo_right_html}
                        {logo_left_html}
                    </div>
                    <div class="brand">تيار الحكمة الوطني</div>
                </header>
                
                <h1>[عنوان رئيسي قوي]</h1>
                <p class="subtitle">[سطر وصفي]</p>
                
                2. Add executive summary:
                <div class="executive-summary">[ملخص من 2-3 فقرات]</div>
                
                3. Add metrics:
                <div class="metrics-grid"><div class="metric-box"><span class="metric-value">XX</span><span class="metric-label">وصف</span></div></div>
                
                4. Use <h2 class="section-title">[عنوان]</h2> for sections
                
                5. Use <div class="callout">...</div> for highlights
                
                6. END with footer
                """

            elif "عرض تقديمي" in report_type:
                target_css = STYLE_PRESENTATION
                file_label = "Presentation"
                
                cover_logos = ""
                if logo_right_html or logo_left_html:
                    cover_logos = f'<div class="cover-logos">{logo_right_html}{logo_left_html}</div>'
                
                header_logos = ""
                if logo_right_html or logo_left_html:
                    header_logos = f'<div class="header-logos">{logo_right_html}{logo_left_html}</div>'
                
                design_rules = f"""
                Create an interactive presentation with slides:
                
                1. FIRST slide (cover):
                <div id="slide-1" class="slide cover active">
                    <div class="cover-content">
                        {cover_logos}
                        <h1 class="main-title">[عنوان العرض]</h1>
                        <p class="sub-title">[العنوان الفرعي]</p>
                    </div>
                </div>
                
                2. Content slides (2, 3, 4...):
                <div id="slide-2" class="slide">
                    <div class="slide-header">
                        <div class="header-title"><h2>[عنوان الشريحة]</h2></div>
                        {header_logos}
                    </div>
                    <div class="slide-content">
                        <div class="text-panel">
                            <h3>[عنوان فرعي]</h3>
                            <p>[محتوى]</p>
                            <ul><li>[نقطة]</li></ul>
                        </div>
                        <div class="visual-panel">
                            <div class="icon-box"><i class="fas fa-chart-line"></i></div>
                            <p>[رقم أو إحصائية]</p>
                        </div>
                    </div>
                </div>
                
                3. Create 5-8 slides based on content
                4. Use FontAwesome icons: fa-chart-line, fa-users, fa-trophy, fa-bullseye, fa-check-circle
                """

            prompt = f"""
            You are an expert Arabic report designer for 'تيار الحكمة الوطني'.
            
            TASK: Create a COMPLETE, DETAILED HTML report.
            
            CRITICAL RULES:
            1. Output ONLY valid HTML body content - NO markdown, NO code blocks
            2. Include ALL data from input - do NOT summarize or skip anything
            3. Follow the design structure EXACTLY as specified
            4. Use proper Arabic formatting (RTL)
            5. Make it visually rich with all specified CSS classes
            
            DESIGN STRUCTURE:
            {design_rules}
            
            INPUT DATA TO PROCESS:
            {full_text}
            
            IMPORTANT: Create a complete, professional report with all sections properly formatted.
            """

            # رسالة التحميل
            status_placeholder = st.empty()
            status_placeholder.markdown('''
            <div class="progress-box">
                <div style="font-size: 2.5rem; margin-bottom: 15px;">🤖</div>
                <div style="color: rgba(255,255,255,0.8); font-size: 1.1rem;">جاري تحليل البيانات وتوليد التقرير...</div>
            </div>
            ''', unsafe_allow_html=True)
            
            try:
                response = model.generate_content(
                    prompt,
                    generation_config=genai.types.GenerationConfig(
                        max_output_tokens=8192,
                        temperature=0.7
                    )
                )
                html_body = clean_html_response(response.text)
            except Exception as api_error:
                status_placeholder.empty()
                st.error(f"❌ خطأ في الاتصال بالذكاء الاصطناعي: {api_error}")
                st.stop()
            
            status_placeholder.empty()
            
            # تجميع HTML النهائي
            container_class = 'presentation-container' if 'عرض تقديمي' in report_type else 'container'
            
            nav_controls = ""
            if 'عرض تقديمي' in report_type:
                nav_controls = """
                <div class="nav-controls">
                    <button class="nav-btn" onclick="prevSlide()"><i class="fas fa-chevron-right"></i></button>
                    <button class="nav-btn" onclick="nextSlide()"><i class="fas fa-chevron-left"></i></button>
                </div>
                <div class="page-number" id="page-num">1 / 1</div>
                """

            final_html = f"""
            <!DOCTYPE html>
            <html lang="ar" dir="rtl">
            <head>
                <meta charset="UTF-8">
                <meta name="viewport" content="width=device-width, initial-scale=1.0">
                <title>تقرير {file_label}</title>
                <link href="https://fonts.googleapis.com/css2?family=Cairo:wght@300;400;600;800&family=Tajawal:wght@400;500;700;800;900&display=swap" rel="stylesheet">
                {target_css}
            </head>
            <body>
                <div class="{container_class}">
                    {html_body}
                    {nav_controls}
                </div>
                {SCRIPT_PRESENTATION if 'عرض تقديمي' in report_type else ''}
            </body>
            </html>
            """
            
            st.session_state['final_html'] = final_html
            st.session_state['file_label'] = file_label

            st.markdown('<div class="success-banner"><span style="color: #22c55e; font-size: 1.2rem; font-weight: 700;">✅ تم إنشاء التقرير بنجاح!</span></div>', unsafe_allow_html=True)
            
            # عرض المعاينة
            st.components.v1.html(final_html, height=800, scrolling=True)
            
            # أزرار التحميل
            st.markdown("<br>", unsafe_allow_html=True)
            
            export_cols = st.columns(3)
            
            with export_cols[0]:
                st.download_button(
                    label="📄 تحميل HTML",
                    data=final_html,
                    file_name=f"{file_label}.html",
                    mime="text/html",
                    use_container_width=True
                )
            
            with export_cols[1]:
                if WEASYPRINT_AVAILABLE:
                    pdf_data = create_pdf(final_html)
                    if pdf_data:
                        st.download_button(
                            label="📕 تحميل PDF",
                            data=pdf_data,
                            file_name=f"{file_label}.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )
                else:
                    st.button("📕 PDF (غير متاح)", disabled=True, use_container_width=True)
            
            with export_cols[2]:
                if DOCX_AVAILABLE:
                    docx_data = create_docx(final_html, file_label)
                    if docx_data:
                        st.download_button(
                            label="📘 تحميل Word",
                            data=docx_data,
                            file_name=f"{file_label}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                else:
                    st.button("📘 Word (غير متاح)", disabled=True, use_container_width=True)

        except Exception as e:
            st.error(f"❌ حدث خطأ: {e}")

# الفوتر
st.markdown("<br><br>", unsafe_allow_html=True)
st.markdown('''
<div style="
    background: linear-gradient(135deg, rgba(26, 54, 93, 0.95), rgba(45, 74, 113, 0.9));
    border-radius: 15px;
    padding: 30px 20px;
    margin: 20px;
    border: 1px solid rgba(212, 175, 55, 0.3);
    text-align: center;
">
    <p style="color: #d4af37; font-size: 1.1rem; font-weight: 700; margin-bottom: 8px;">الجهاز المركزي للجودة الشاملة</p>
    <p style="color: rgba(255, 255, 255, 0.8); font-size: 1rem;">وحدة التخطيط الاستراتيجي والتطوير</p>
    <p style="color: rgba(255, 255, 255, 0.5); font-size: 0.85rem; margin-top: 15px;">جميع الحقوق محفوظة © 2026</p>
</div>
''', unsafe_allow_html=True)
